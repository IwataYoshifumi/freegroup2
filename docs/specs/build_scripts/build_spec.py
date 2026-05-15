"""build_spec.py — FreeGroup2 名刺画像取り込みOCR仕様書 v1.3.4 生成スクリプト

【使い方】
    1. 依存解決
        python -m venv .venv
        .venv\\Scripts\\activate              # Windows
        # source .venv/bin/activate          # macOS / Linux
        pip install -r requirements.txt
    2. 実行
        python build_spec.py
    3. 出力
        ../v1_3_4/名刺画像取り込みOCR仕様書_v1_3_4.docx に保存される

【方針】
    - HOW（コード詳細）はコード（GitHub）が Single Source of Truth、本仕様書は WHAT / WHY を記述
    - 概念図は ./diagrams/ 配下の PNG をそのまま埋め込む（ソース図形は別途管理）
    - 次バージョン（v1.4.0 以降）は本ファイルをコピーして差分修正する想定

【スタイル要件】
    - フォント：Yu Gothic
    - 本文：Pt(11)、表内テキスト：Pt(10)
    - H1 #1F4E79 Pt(16) / H2 #2E75B6 Pt(13) / H3 #5B9BD5 Pt(12) / H4 #C00000 Pt(11)（すべて太字）
    - 表ヘッダー背景 #DEEAF6、★改訂セル背景 #FFF2CC
    - A4 縦、上下左右マージン 1.9cm
    - フッターにページ番号（X / Y、中央揃え）
    - 表は cantSplit、見出しは keep_with_next でページまたぎ防止
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# ====================================================================
# 設定
# ====================================================================
SCRIPT_DIR = Path(__file__).resolve().parent
DIAGRAMS_DIR = SCRIPT_DIR / "diagrams"
OUTPUT_DIR = SCRIPT_DIR.parent / "v1_3_4"
OUTPUT_FILE = OUTPUT_DIR / "名刺画像取り込みOCR仕様書_v1_3_4.docx"

FONT_JP = "Yu Gothic"
COLOR_H1 = RGBColor(0x1F, 0x4E, 0x79)
COLOR_H2 = RGBColor(0x2E, 0x75, 0xB6)
COLOR_H3 = RGBColor(0x5B, 0x9B, 0xD5)
COLOR_H4 = RGBColor(0xC0, 0x00, 0x00)
COLOR_GRAY = RGBColor(0x59, 0x59, 0x59)
COLOR_HEADER_BG = "DEEAF6"
COLOR_REVISION_BG = "FFF2CC"

REV_MARK = "★ "  # ★改訂表記


# ====================================================================
# 低レベルユーティリティ
# ====================================================================
def _set_run_font(run, size_pt: float, *, bold: bool = False, color: RGBColor | None = None,
                  italic: bool = False) -> None:
    run.font.name = FONT_JP
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_JP)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _set_cell_bg(cell, color_hex: str) -> None:
    """セルに塗り色を適用する（既存 shd は除去してから上書き）。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    # 既存 shd 要素があれば削除してから新規追加（重複と他スタイル干渉を防止）
    for existing in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _set_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def _set_keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def _setup_page(doc) -> None:
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(1.9)
    section.bottom_margin = Cm(1.9)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)


def _add_page_number_footer(doc) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _field(run_text: str):
        run = p.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = run_text
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_end)
        _set_run_font(run, 9)
        return run

    _field("PAGE")
    sep = p.add_run(" / ")
    _set_run_font(sep, 9)
    _field("NUMPAGES")


# ====================================================================
# 段落・見出し
# ====================================================================
def add_paragraph(doc, text: str, *, size: float = 11, bold: bool = False,
                  color: RGBColor | None = None, italic: bool = False,
                  alignment=None, keep_with_next: bool = False):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    if keep_with_next:
        _set_keep_with_next(p)
    run = p.add_run(text)
    _set_run_font(run, size, bold=bold, color=color, italic=italic)
    return p


def _add_heading(doc, text: str, *, style_name: str, size_pt: float,
                 color: RGBColor, space_before: float, space_after: float):
    """Heading n スタイル付きの見出し段落を追加し、フォントを上書きする。

    Word のナビゲーションウィンドウ・目次自動生成に使う pStyle="Heading n" を持たせ、
    ランレベルでフォント名・サイズ・色・太字を上書きする（テーマ色を無効化）。
    """
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    _set_keep_with_next(p)
    run = p.add_run(text)
    _set_run_font(run, size_pt, bold=True, color=color)
    # Heading スタイルが持つ themeColor を無効化（color.rgb で上書き済みだが念のため）
    rpr = run._element.get_or_add_rPr()
    for tag in ("w:color",):
        for el in rpr.findall(qn(tag)):
            if el is not run.font.color._element:  # noqa: SLF001
                continue
    return p


def add_h1(doc, text: str, *, mark_revision: bool = False):
    if mark_revision:
        text = REV_MARK + text
    return _add_heading(doc, text, style_name="Heading 1", size_pt=16,
                        color=COLOR_H1, space_before=18, space_after=8)


def add_h2(doc, text: str, *, mark_revision: bool = False):
    if mark_revision:
        text = REV_MARK + text
    return _add_heading(doc, text, style_name="Heading 2", size_pt=13,
                        color=COLOR_H2, space_before=12, space_after=6)


def add_h3(doc, text: str, *, mark_revision: bool = False):
    if mark_revision:
        text = REV_MARK + text
    return _add_heading(doc, text, style_name="Heading 3", size_pt=12,
                        color=COLOR_H3, space_before=8, space_after=4)


def add_h4(doc, text: str, *, mark_revision: bool = False):
    if mark_revision:
        text = REV_MARK + text
    return _add_heading(doc, text, style_name="Heading 4", size_pt=11,
                        color=COLOR_H4, space_before=6, space_after=3)


def add_bullet(doc, text: str, *, indent_cm: float = 0.5, size: float = 11):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent_cm + 0.6)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    run = p.add_run("・" + text)
    _set_run_font(run, size)
    return p


def add_bullets(doc, items: list[str]) -> None:
    for item in items:
        add_bullet(doc, item)


# ====================================================================
# テーブル
# ====================================================================
def add_table(doc, headers: list[str], rows: list[list[str]], *,
              col_widths_cm: list[float] | None = None,
              revision_row_indices: set[int] | None = None,
              revision_col_indices: set[int] | None = None) -> None:
    """ヘッダー＋本体表を追加。

    revision_row_indices: 行インデックス（0始まり）の集合。指定行は黄背景。
    revision_col_indices: 列インデックス（0始まり）の集合。指定セルは黄背景。
    """
    revision_row_indices = revision_row_indices or set()
    revision_col_indices = revision_col_indices or set()

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ヘッダー
    header_row = table.rows[0]
    _set_cant_split(header_row)
    for i, htext in enumerate(headers):
        cell = header_row.cells[i]
        _set_cell_bg(cell, COLOR_HEADER_BG)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(htext)
        _set_run_font(run, 10, bold=True)

    # 本体
    for r_idx, row_data in enumerate(rows):
        body_row = table.rows[1 + r_idx]
        _set_cant_split(body_row)
        rev_row = r_idx in revision_row_indices
        for c_idx, ctext in enumerate(row_data):
            cell = body_row.cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if rev_row or c_idx in revision_col_indices:
                _set_cell_bg(cell, COLOR_REVISION_BG)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(ctext))
            _set_run_font(run, 10)

    # カラム幅
    if col_widths_cm is not None and len(col_widths_cm) == len(headers):
        for row in table.rows:
            for i, w in enumerate(col_widths_cm):
                row.cells[i].width = Cm(w)

    # 表の後に少し空白
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


# ====================================================================
# 画像
# ====================================================================
def add_image(doc, png_filename: str, caption: str, *, width_inches: float = 6.0) -> None:
    img_path = DIAGRAMS_DIR / png_filename
    if not img_path.exists():
        raise FileNotFoundError(f"diagram not found: {img_path}")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_keep_with_next(p)
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(width_inches))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    _set_run_font(cap_run, 9, italic=True, color=COLOR_GRAY)


# ====================================================================
# タイトルページ
# ====================================================================
def add_title_page(doc) -> None:
    # 上部余白
    for _ in range(6):
        doc.add_paragraph()

    add_paragraph(doc, "名刺画像取り込みOCR 仕様書", size=24, bold=True,
                  color=COLOR_H1, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "バージョン v1.3.4", size=18, bold=True,
                  color=COLOR_H2, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_paragraph(doc, "FreeGroup2 名刺管理機能", size=14,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "2026年4月作成 / 5月1日改訂（見栄え修正パッチ）",
                  size=11, color=COLOR_GRAY,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # 改ページ
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ====================================================================
# 本文セクション
# ====================================================================
def section_revision_summary(doc) -> None:
    add_h1(doc, "v1.3.4 改訂サマリー")
    add_paragraph(doc,
        "本書は v1.3.3 の見栄え修正パッチリリースである。仕様書本文のコンテンツに変更はなく、"
        "Word での閲覧性を高めるための 3 点（見出しスタイル・表ヘッダー背景・★改訂セル背景）を"
        "整備した。HOW は GitHub の実コードを Single Source of Truth とする方針は維持する。"
    )
    doc.add_paragraph()
    add_paragraph(doc,
        "本仕様書の方針：v1.3.4 は完結版として、v1.0.0〜v1.3.3 の累積仕様を本書一冊で参照可能な形に統合する。"
        "コードの正解は実装ファイル（GitHub）を Single Source of Truth とし、"
        "仕様書は「何を」「なぜ」を記述する役割に専念する。"
        "★マーク は v1.3.x で改訂した箇所、または堅牢化関連の重要記述に付与する。"
    )

    add_table(doc,
        headers=["#", "修正内容"],
        rows=[
            ["1", "見出しスタイルに Word の Heading 1〜4 を適用し、ナビゲーションウィンドウと目次自動生成に対応"],
            ["2", "表ヘッダー行に薄青背景 #DEEAF6 を適用（既存の塗りを堅牢化、表スタイル干渉を排除）"],
            ["3", "★改訂セルに黄色背景 #FFF2CC を適用し、v1.3.x で改訂・新規追加された行を視認性高く表示"],
        ],
        col_widths_cm=[1.2, 15.5],
        revision_row_indices={0, 1, 2},
    )


def section_chapter1(doc) -> None:
    add_h1(doc, "第1章 はじめに")
    add_h2(doc, "1.1 目的")
    add_paragraph(doc,
        "本仕様書は、FreeGroup2 における名刺画像取り込み・OCR・データ管理機能の設計を定義することを"
        "目的とする。本機能は、ユーザーがアップロードした名刺画像から名刺情報を自動抽出し、"
        "データベースに保存・管理するシステムである。"
    )
    add_h2(doc, "1.2 適用範囲")
    add_paragraph(doc, "本仕様書は FreeGroup2 の v1.3.4 における名刺管理機能を対象とし、以下の機能を含む。")
    add_bullets(doc, [
        "画像アップロード機能",
        "名刺領域の自動検出機能（OpenCV による事前検出）",
        "名刺画像の正規化機能（透視変換・縦書き対応）",
        "OCR 実行機能（Claude Sonnet 4.6 による）",
        "OCR 結果の JSON 保存機能",
        "Person / Contact データの管理機能",
        "信頼度メタデータの管理機能",
        "OCR パイプラインの堅牢性管理機能（多重起動対策・stuck 検出・整合性検査）",
    ])
    add_h2(doc, "1.3 用語定義")
    add_table(doc,
        headers=["用語", "定義"],
        rows=[
            ["元画像", "ユーザーがアップロードした写真。複数枚の名刺を含む可能性がある"],
            ["名刺画像", "元画像から個別の名刺領域を切り抜いて正規化した画像"],
            ["bbox", "Bounding Box。画像内の物体の位置を表す矩形"],
            ["raw_json", "OCR バックエンドから返される生の JSON 結果。不変データとして保管する"],
            ["Tool Use", "Claude API の構造化出力機能。JSON Schema を指定して安定した形式で結果を取得する"],
            ["Person", "人物の本体。複数の Contact が紐付く可能性がある"],
            ["Contact", "名刺ごとのスナップショット。名前・会社・連絡先などのフルセットを保持する"],
            ["ContactFieldConfidence", "Contact フィールドの信頼度メタデータ"],
            ["worker", "OCR 処理を実行する管理コマンド（process_pending）のプロセス"],
            ["CAS", "Compare-And-Swap。楽観ロック方式。ある状態のときのみ別の状態に遷移させる原子的操作"],
            ["stuck sweeper", "一定時間以上 processing のままのレコードを救済する仕組み（30 分しきい値）"],
            ["純関数", "DB を一切触らない、副作用なし、同じ入力で同じ出力の関数"],
            ["準関数", "DB を読むが書かない、外部世界に副作用なしの関数"],
            ["副作用あり関数", "DB 書き込み・例外送出・API 呼び出し・ファイル書き込みなどを行う関数"],
        ],
        col_widths_cm=[4.5, 12.0],
    )


def section_chapter2(doc) -> None:
    add_h1(doc, "第2章 システム概要")
    add_h2(doc, "2.1 機能概要")
    add_paragraph(doc,
        "本機能は、ユーザーが撮影した名刺画像を取り込み、AI（Claude API）を活用して名刺情報を自動抽出し、"
        "構造化されたデータとしてデータベースに保存する。1 枚の画像に複数の名刺が含まれていても、"
        "OpenCV による事前検出と Claude による個別 OCR を組み合わせて処理する。"
    )
    add_h2(doc, "2.2 採用する処理方式")
    add_paragraph(doc,
        "v1.3.x では、画像内の名刺検出を OpenCV で行い、各名刺の OCR を Claude API で行う方式を採用する。"
        "横長統一処理は v1.3.0 で削除され、縦書き名刺にも対応している。"
    )
    add_h2(doc, "2.3 処理フローの全体像")
    add_paragraph(doc,
        "システム全体の処理フローは以下のとおり。ユーザーの画像アップロード時には View 層が同期的に"
        "元画像を保存し、OCR 処理は cron 起動の Task 層が非同期で実行する。"
    )
    add_image(doc, "A_flow.png", "図 2-1：システム全体の処理フロー（View 層と Task 層の責務分離）")
    add_paragraph(doc,
        "View 層の責務は元画像の保存（OriginalImage 作成、status=pending）までに限定される。"
        "OCR 処理は別プロセス（cron 経由の process_pending）が pending レコードを拾って実行する。"
        "この設計により、ユーザーは長時間の OCR 処理を待たずに画面を遷移できる。"
    )


def section_chapter3(doc) -> None:
    add_h1(doc, "第3章 機能要件")
    add_h2(doc, "3.1 画像アップロード機能")
    add_h3(doc, "3.1.1 受け付ける画像形式")
    add_bullets(doc, [
        "JPEG (.jpg, .jpeg)",
        "PNG (.png)",
    ])
    add_paragraph(doc, "HEIC は v1.3.x 時点では受付不可（v1.4.0 以降で対応検討）")
    add_h3(doc, "3.1.2 バリデーション仕様")
    add_bullets(doc, [
        "最大ファイルサイズ：5MB",
        "受付拡張子：jpg / jpeg / png のみ",
        "最低画像サイズ：100 × 100px（OpenCV 処理上の最低基準）",
    ])
    add_h2(doc, "3.2 OCR 処理の起動", mark_revision=True)
    add_paragraph(doc,
        "OCR 処理を開始するユーザーアクションは「画像アップロード」のみとする。"
        "失敗した OriginalImage に対する retry 機能はユーザーに提供しない。"
    )
    add_paragraph(doc, "【ユーザーが失敗を確認した場合の対処】", bold=True)
    add_bullets(doc, [
        "画像を撮り直して新しくアップロードする",
        "不要な失敗 OriginalImage は削除する（削除機能は Phase 4 で実装）",
    ])
    add_paragraph(doc, "【retry_failed_ocr 管理コマンドの位置付け】", bold=True)
    add_bullets(doc, [
        "開発・運用ツールとして提供する",
        "開発時：プロンプト修正後の動作確認用",
        "運用時：API 一時障害後のまとめリトライ用",
        "本番運用ではエンドユーザーに再アップロードを促す",
    ])


def section_chapter4(doc) -> None:
    add_h1(doc, "第4章 データベース設計")
    add_h2(doc, "4.1 全体構造")
    add_paragraph(doc,
        "v1.3.4 のデータモデル全体構造を以下に示す。OriginalImage を起点として、"
        "BusinessCard・Contact・Person・ContactFieldConfidence が階層的に紐付く。"
    )
    add_image(doc, "B_er.png", "図 4-1：データモデルの ER 図")

    # 4.2 OriginalImage
    add_h2(doc, "4.2 OriginalImage（元画像DB）", mark_revision=True)
    add_paragraph(doc, "OriginalImage は元画像 1 件に対応するレコード。")
    add_table(doc,
        headers=["フィールド名", "型", "説明"],
        rows=[
            ["id", "UUIDField (PK)", "プライマリキー（uuid.uuid4）"],
            ["user", "FK(User, CASCADE)", "アップロードしたユーザー"],
            ["image_file", "ImageField", "元画像ファイル（originals/%Y/%m/%d/）"],
            ["status", "CharField(20)", "処理状態（5 値、4.7.1 参照）"],
            ["claimed_at", "DateTimeField (nullable)", "CAS で processing 遷移時刻を記録。stuck sweeper 判定用（v1.3.0 新規）"],
            ["raw_json", "JSONField (nullable)", "OCR 結果 JSON（不変）"],
            ["detected_count", "IntegerField (default=0)", "検出された名刺数"],
            ["error_message", "TextField (default=\"\")", "失敗理由・部分失敗ログを集約"],
            ["created_at", "DateTimeField (auto_now_add)", "レコード作成日時"],
            ["updated_at", "DateTimeField (auto_now)", "最終更新日時"],
        ],
        col_widths_cm=[4.0, 4.5, 8.0],
        revision_row_indices={4},
    )
    add_paragraph(doc, "【raw_json への保存制約】", bold=True)
    add_bullets(doc, [
        "raw_json には OCR 結果 JSON（第5章で定義する標準形式）のみを保存",
        "画像バイナリ・base64 エンコードされた画像は保存しない",
        "理由：JSONField の肥大化を防ぐため。画像本体はファイルシステムで管理",
    ])
    add_paragraph(doc, "【STATUS_CHOICES】", bold=True)
    add_table(doc,
        headers=["定数名", "コード値", "表示名"],
        rows=[
            ["STATUS_PENDING", '"pending"', "処理待ち"],
            ["STATUS_PROCESSING", '"processing"', "処理中（v1.3.0 新規）"],
            ["STATUS_EXTRACTED", '"extracted"', "完了"],
            ["STATUS_GARBAGE", '"garbage"', "無効画像"],
            ["STATUS_FAILED", '"failed"', "処理失敗"],
        ],
        col_widths_cm=[5.0, 4.0, 7.5],
        revision_row_indices={1},
    )

    # 4.3 BusinessCard
    add_h2(doc, "4.3 BusinessCard（名刺DB）", mark_revision=True)
    add_table(doc,
        headers=["フィールド名", "型", "説明"],
        rows=[
            ["id", "UUIDField (PK)", "プライマリキー"],
            ["original_image", "FK(OriginalImage, CASCADE)", "元画像への外部キー"],
            ["card_image", "ImageField (nullable)", "正規化後の名刺画像。切り抜き失敗時は null"],
            ["card_index", "IntegerField", "raw_json[\"cards\"] 配列内インデックス"],
            ["orientation", "CharField(20, choices)", "名刺の向き（5 値、下表参照）"],
            ["created_at", "DateTimeField (auto_now_add)", "レコード作成日時"],
            ["updated_at", "DateTimeField (auto_now)", "最終更新日時"],
        ],
        col_widths_cm=[4.0, 4.5, 8.0],
    )
    add_paragraph(doc, "【ORIENTATION_CHOICES】（v1.3.0 で 5 値に整理）", bold=True)
    add_table(doc,
        headers=["定数名", "コード値", "表示名"],
        rows=[
            ["ORIENTATION_NORMAL", '"normal"', "正位（既定値）"],
            ["ORIENTATION_90_CW", '"rotate_90_cw"', "時計回り 90°"],
            ["ORIENTATION_90_CCW", '"rotate_90_ccw"', "反時計回り 90°"],
            ["ORIENTATION_180", '"rotate_180"', "180° 回転"],
            ["ORIENTATION_MIRROR", '"mirror"', "鏡像"],
        ],
        col_widths_cm=[5.5, 4.5, 6.5],
        revision_row_indices={0, 1, 2, 3, 4},
    )
    add_paragraph(doc, "【制約】", bold=True)
    add_bullets(doc, [
        "UniqueConstraint: (original_image, card_index) の組み合わせは一意",
        "on_delete=CASCADE: OriginalImage 削除で連鎖削除",
    ])

    # 4.4 Contact
    add_h2(doc, "4.4 Contact（連絡先DB）", mark_revision=True)
    add_paragraph(doc,
        "名刺ごとのスナップショット。BusinessCard と OneToOne 関係。Person への FK は NOT NULL。"
    )
    add_paragraph(doc, "【主要フィールド】", bold=True)
    add_table(doc,
        headers=["フィールド名", "型", "説明"],
        rows=[
            ["id", "UUIDField (PK)", "プライマリキー"],
            ["business_card", "OneToOne(BusinessCard, CASCADE)", "名刺との 1 対 1 関係"],
            ["person", "FK(Person, CASCADE)", "人物本体への参照（NOT NULL）"],
            ["full_name", "CharField(255)", "氏名（フル）"],
            ["last_name", "CharField(255)", "姓"],
            ["first_name", "CharField(255)", "名"],
            ["salutation_name", "CharField(255)", "敬称表記"],
            ["company", "CharField(255)", "会社名"],
            ["department / title", "CharField(255)", "部署 / 役職"],
            ["qualification / catchphrase", "CharField(500)", "資格 / キャッチフレーズ"],
            ["branch", "CharField(255)", "支店・営業所"],
            ["address", "CharField(500)", "住所"],
            ["email", "CharField(255)", "メールアドレス"],
            ["phone / mobile / fax", "CharField(50)", "固定電話 / 携帯 / FAX"],
            ["website", "CharField(500)", "ウェブサイト URL"],
            ["twitter / instagram / github", "CharField(255)", "SNS（短文系）"],
            ["linkedin / facebook", "CharField(500)", "SNS（URL 系）"],
            ["notes", "TextField", "自由記述メモ"],
            ["created_at / updated_at", "DateTimeField", "自動付与（auto_now_add / auto_now）"],
        ],
        col_widths_cm=[4.5, 4.5, 7.5],
    )
    add_paragraph(doc, "値が無い場合は空文字（default=\"\"）。null は使わない方針。")

    # 4.5 Person
    add_h2(doc, "4.5 Person（人物DB）")
    add_paragraph(doc,
        "人物の本体。複数の Contact が紐付く可能性がある（同じ人の名刺を複数枚もらった場合）。"
        "v1.3.4 時点では Contact 作成時に同時生成され、自動マージは行わない（v2.0.0 で検討）。"
    )
    add_paragraph(doc, "【主要フィールド】", bold=True)
    add_bullets(doc, [
        "id: UUIDField (PK)",
        "created_at / updated_at: DateTimeField",
    ])

    # 4.6 ContactFieldConfidence
    add_h2(doc, "4.6 ContactFieldConfidence（信頼度メタDB）", mark_revision=True)
    add_paragraph(doc,
        "Contact フィールドごとの信頼度を別テーブルで管理する。"
        "human-in-the-loop による確認履歴も保持する。"
    )
    add_paragraph(doc, "【主要フィールド】", bold=True)
    add_table(doc,
        headers=["フィールド名", "型", "説明"],
        rows=[
            ["id", "UUIDField (PK)", "プライマリキー"],
            ["contact", "FK(Contact, CASCADE)", "related_name=\"confidences\""],
            ["field_name", "CharField(50)", "Contact 側のフィールド名（例: \"email\", \"phone\"）"],
            ["confidence", "CharField(10, choices)", "low / medium のみ（high は記録対象外）"],
            ["confirmed_at", "DateTimeField (nullable)", "ユーザーが確認した日時（human-in-the-loop 用）"],
            ["confirmed_by", "FK(User, SET_NULL)", "確認したユーザー（human-in-the-loop 用）"],
            ["created_at / updated_at", "DateTimeField", "自動付与"],
        ],
        col_widths_cm=[4.5, 4.5, 7.5],
        revision_row_indices={4, 5},
    )
    add_paragraph(doc, "【方針】", bold=True)
    add_bullets(doc, [
        "high の値はレコードを作成しない（信頼できる値はメタ管理不要）",
        "medium / low の値のみ ContactFieldConfidence に記録",
        "UI で「要確認フィールド」として表示する根拠データ",
        "人による確認後、confirmed_at / confirmed_by を更新",
        "UniqueConstraint: (contact, field_name) の組み合わせは一意",
    ])

    # 4.7 ステータス定義
    add_h2(doc, "4.7 ステータス定義")
    add_h3(doc, "4.7.1 OriginalImage.status の判定ルール")
    add_paragraph(doc,
        "状態遷移は以下の図に示すとおり。v1.3.0 で processing 状態を追加し、CAS 楽観ロックの土台とした。"
        "stuck sweeper による pending 差し戻しは例外経路として位置付ける（8.5.1.1 参照）。"
    )
    add_image(doc, "C_status.png", "図 4-2：OriginalImage.status の遷移図")
    # ★修正④：garbage 行を実コードに合わせて修正
    add_table(doc,
        headers=["値", "日本語表示", "判定条件"],
        rows=[
            ["pending", "処理待ち", "アップロード直後〜tasks 層の処理開始前"],
            ["processing", "処理中", "CAS で claim_lock 成功〜run_pipeline 完了前。stuck sweeper の対象"],
            ["extracted", "完了", "名刺検出 ≧1 件。has_minimum_info で全 card 落ちした場合も BusinessCard 0 件で extracted に遷移する（v1.3.3 改訂）"],
            ["garbage", "無効画像", "名刺検出 0 件（detect_cards 結果が空）"],
            ["failed", "処理失敗", "API 通信失敗 / スキーマ検証失敗 / DB 書き込み失敗 / 予期しない例外"],
        ],
        col_widths_cm=[3.0, 3.0, 10.5],
        revision_row_indices={2, 3},
    )
    add_h3(doc, "4.7.2 ContactFieldConfidence.confidence")
    add_bullets(doc, [
        "high：信頼度高（人による確認不要）。レコードを作成しない（DB 保存対象外）",
        "medium：信頼度中（要確認推奨）",
        "low：信頼度低（必ず確認）",
    ])

    # 4.8 リレーション削除ルール
    add_h2(doc, "4.8 リレーション削除ルール")
    add_bullets(doc, [
        "OriginalImage.delete() → BusinessCard / Contact / ContactFieldConfidence が CASCADE 削除",
        "BusinessCard.delete() → Contact / ContactFieldConfidence が CASCADE 削除",
        "Contact.delete() → ContactFieldConfidence が CASCADE 削除（Person は残る）",
        "Person.delete() → 紐付く Contact が CASCADE 削除（Contact.person が NOT NULL の双方向 CASCADE）",
    ])


def section_chapter5(doc) -> None:
    add_h1(doc, "第5章 OCR 結果 JSON の仕様")
    add_h2(doc, "5.1 JSON 構造の概要")
    add_paragraph(doc,
        "OCR バックエンドから返される標準形式の JSON 構造。"
        "OriginalImage.raw_json にそのまま保存される。"
    )
    add_paragraph(doc, "【トップレベル構造】", bold=True)
    add_bullets(doc, [
        "schema_version：必須。例: \"1.3.0\"",
        "ocr_meta：OCR 実行情報（timestamp 等）",
        "api_response：API 応答メタ（model / stop_reason / usage 等。実コードで保存される）",
        "cards：名刺ごとの配列。各要素は card_index / is_business_card / orientation / fields / fields_array を持つ",
    ])
    # ★修正③：JSON Schema ファイル名
    add_h2(doc, "5.2 schema_version の管理", mark_revision=True)
    add_bullets(doc, [
        "schema_version は標準 JSON の必須フィールド",
        "フォーマット定義は docs/json_schema/v{X.Y.Z}/combined_response.json として Git 管理（v1.3.3 改訂）",
        "v1.3.x 時点での標準フォーマットは schema_version=\"1.3.0\"",
        "json_normalizer は schema_version を確認して適切な解析ロジックを適用",
    ])
    add_h2(doc, "5.3 raw_json は不変・card_index は不変")
    add_bullets(doc, [
        "OriginalImage.raw_json は OCR 時点の証跡として保存し、その後変更してはならない",
        "cards 配列の順序は raw_json 保存時点で固定される。差し替え禁止",
        "BusinessCard.card_index は raw_json[\"cards\"] 配列内の固定インデックスを参照",
        "再 OCR を行う場合は raw_json を上書きせず、新しい再処理設計で扱う（v2.0.0 で検討）",
    ])
    add_h2(doc, "5.4 confidence 値の扱い")
    add_bullets(doc, [
        "high / medium / low の 3 値のみ",
        "orientation 別の自動調整あり（calc_orientation_adjusted_confidence_map 純関数）",
        "縦書き名刺（rotate_90_cw / rotate_90_ccw）は横書きより一般的に confidence が低めに調整される",
    ])


def section_chapter6(doc) -> None:
    add_h1(doc, "第6章 画像処理仕様")
    add_h2(doc, "6.1 名刺検出（OpenCV）")
    add_paragraph(doc,
        "元画像から名刺領域を検出する処理は OpenCV で実装する。"
        "輝度差・Canny エッジ・HSV 彩度の 3 種マスクを OR 合成して候補を抽出し、"
        "minAreaRect で 4 隅座標を取得、透視変換で正立化した画像を返す（実装は opencv_detector.py を参照）。"
    )
    add_paragraph(doc, "【処理ステップ概要】", bold=True)
    add_bullets(doc, [
        "グレースケール変換 / クローズ処理によるノイズ除去",
        "3 種マスク（輝度差 / Canny エッジ / HSV 彩度）の OR 合成",
        "輪郭検出（findContours）",
        "minAreaRect → boxPoints で矩形化",
        "面積・アスペクト比・重複でフィルタリング",
        "top_left の y 優先・x 次でソート（card_index の順序確定）",
        "透視変換（getPerspectiveTransform → warpPerspective）",
    ])
    add_paragraph(doc, "【v1.3.0 改訂】", bold=True)
    add_bullets(doc, [
        "横長統一処理を削除（縦書き名刺対応）",
        "縦書き名刺はそのまま縦向きで Claude に渡す",
        "orientation 判定は Claude 側で行う",
    ])
    # ★修正⑤：最低画像サイズ判定の所在
    add_h2(doc, "6.2 切り抜き失敗時の扱い", mark_revision=True)
    add_bullets(doc, [
        "切り抜き後の画像 width < 100 または height < 50 は失敗扱い（判定は opencv_detector._warp_card() 内で行う。CardCropper には到達しない）",
        "切り抜き失敗時も BusinessCard / Contact は作成可能（card_image=null）",
        "失敗理由は OriginalImage.error_message に記録",
    ])


def section_chapter7(doc) -> None:
    add_h1(doc, "第7章 OCR バックエンド仕様")
    add_h2(doc, "7.1 サポートする OCR バックエンド")
    add_table(doc,
        headers=["バックエンド", "v1.3.x 状況", "備考"],
        rows=[
            ["Claude Sonnet 4.6", "採用中", "Tool Use による構造化出力。is_business_card / orientation / フィールド抽出を統合実行"],
            ["GPT-4o", "候補", "Structured Output 対応。v2.0.0 で切替候補"],
            ["Google Cloud Vision", "候補", "テキスト抽出 + LLM 構造化の 2 段階処理が必要"],
            ["EasyOCR", "候補", "ローカル動作。プライバシー要件対応"],
        ],
        col_widths_cm=[4.0, 3.0, 9.5],
    )
    add_h2(doc, "7.2 出力形式の統一方針")
    add_paragraph(doc,
        "すべての OCR バックエンドは、第5章で定義する標準 JSON 形式（schema_version / ocr_meta / api_response / cards）"
        "で結果を返す。"
    )
    add_paragraph(doc, "【方針1：構造化出力対応バックエンド】", bold=True)
    add_paragraph(doc,
        "Claude Tool Use / GPT-4o Structured Output 等。"
        "プロンプトおよび JSON Schema 指定により、標準形式を直接出力させる。"
    )
    add_paragraph(doc, "【方針2：構造化出力非対応バックエンド】", bold=True)
    add_paragraph(doc,
        "Google Cloud Vision / EasyOCR 等。"
        "バックエンド内部で「テキスト抽出 → LLM で構造化」の 2 段階処理を行う。"
    )
    add_h2(doc, "7.3 プロンプト仕様")
    add_paragraph(doc,
        "Claude へのプロンプトは cards/prompts/extract_combined.txt に配置。v1.3.0 で以下の改善を実施。"
    )
    add_bullets(doc, [
        "ハルシネーション抑制（推測フィールドを生成しない指示を強化）",
        "department / title 分離精度向上",
        "qualification / catchphrase / branch フィールド追加",
    ])
    # ★修正③：JSON Schema ファイル名
    add_h2(doc, "7.4 JSON Schema ファイルの管理", mark_revision=True)
    add_bullets(doc, [
        "docs/json_schema/v{X.Y.Z}/combined_response.json で Git 管理（v1.3.3 改訂）",
        "Claude Tool Use の input_schema にそのまま指定（API レベルでスキーマ強制）",
        "受信レスポンスの jsonschema 検証（スキーマ違反検出）",
        "仕様書とコードの「契約」を一元化（Single Source of Truth）",
    ])


def section_chapter8(doc) -> None:
    add_h1(doc, "第8章 Django 実装設計", mark_revision=True)
    add_h2(doc, "8.1 アプリケーション構成")
    add_paragraph(doc, "【ディレクトリ構成】", bold=True)
    add_bullets(doc, [
        "cards/views.py — View 層",
        "cards/urls.py / cards/originals_urls.py — URL ルーティング",
        "cards/forms.py — UploadForm（validate_image を clean_image() で呼ぶ）",
        "cards/models.py — Model 層",
        "cards/services/ — Service 層（純関数中心）",
        "cards/services/detectors/ — 名刺検出バックエンド（opencv_detector 等）",
        "cards/tasks/ — Task 層（副作用あり処理）",
        "cards/management/commands/ — 管理コマンド",
        "cards/prompts/ — Claude プロンプト（extract_combined.txt）",
        "cards/migrations/ — DB マイグレーション",
        "docs/json_schema/v1.3.0/combined_response.json — JSON Schema 定義",
    ])

    # 8.2 ルーティング
    add_h2(doc, "8.2 ルーティング（URL パターン）", mark_revision=True)
    add_table(doc,
        headers=["URL", "namespace:name", "View"],
        rows=[
            ["/", "home", "HomeView（クラス View、Phase 4 で home アプリに移行）"],
            ["/cards/upload/", "cards:card_upload", "UploadView"],
            ["/cards/", "cards:card_list", "CardListView"],
            ["/cards/<uuid:pk>/", "cards:card_detail", "CardDetailView"],
            ["/originals/", "originals:original_list", "OriginalListView"],
            ["/originals/<uuid:pk>/", "originals:original_detail", "OriginalDetailView"],
        ],
        col_widths_cm=[5.5, 5.0, 6.0],
    )

    # 8.3 View 層 ★修正①②
    add_h2(doc, "8.3 View 層の設計", mark_revision=True)
    add_paragraph(doc,
        "View 層は薄く保ち、ビジネスロジックは services / tasks 層に委譲する。"
        "v1.3.x 時点では認証は仮実装とし、cards.views.get_current_user() が "
        "「request.user が認証済みならそのまま、未認証ならスーパーユーザーフォールバック」を返す。"
        "Phase 4 で LoginRequiredMixin による本格認証を導入予定（§11.4 / §12.1 参照）。"
    )
    add_paragraph(doc, "【主要 View 一覧】", bold=True)
    add_table(doc,
        headers=["View 名", "継承", "用途"],
        rows=[
            ["HomeView", "TemplateView + LoginRequiredMixin", "ホーム画面表示（Phase 4 で home アプリに移行、紐付けアラート含む）"],
            ["UploadView", "FormView", "名刺画像アップロード受付・OriginalImage 作成（status=pending で保存）"],
            ["CardListView", "ListView", "名刺一覧・7 フィールド AND 検索"],
            ["CardDetailView", "DetailView", "名刺詳細・Contact / ContactFieldConfidence 表示"],
            ["OriginalListView", "ListView", "元画像一覧・status / 日付フィルタ"],
            ["OriginalDetailView", "DetailView", "元画像詳細・紐付き BusinessCard 一覧"],
        ],
        col_widths_cm=[4.5, 3.0, 9.0],
    )
    add_paragraph(doc, "【UploadView の責務（v1.3.3 修正）】", bold=True)
    add_bullets(doc, [
        "image_processor.validate_image でファイル形式・サイズをバリデーション（実装は UploadForm.clean_image() 経由で呼ばれる）",
        "image_processor.convert_to_jpeg で純粋 JPEG（EXIF orientation 適用済み）に変換",
        "OriginalImage を status=pending で保存（同期処理）",
        "OCR 処理は行わない。process_pending（cron）が後で拾って実行する",
        "保存後、アップロードした元画像の詳細画面（originals:original_detail）にリダイレクト（v1.3.3 改訂）",
    ])

    # 8.4 Service 層
    add_h2(doc, "8.4 Service 層の設計（cards/services/）", mark_revision=True)
    add_paragraph(doc,
        "Service 層は純関数中心で構成し、DB 操作・外部 API 呼び出しを行わない。"
        "副作用がある場合は明示的に docstring で示す。"
    )
    add_h3(doc, "8.4.1 image_processor.py")
    add_table(doc,
        headers=["関数名", "性質", "用途"],
        rows=[
            ["validate_image", "副作用あり（例外 raise）", "ファイル形式・サイズ検証"],
            ["convert_to_jpeg", "副作用あり（画像変換）", "画像を JPEG バイト列に変換（EXIF orientation 物理適用）"],
        ],
        col_widths_cm=[4.5, 5.5, 6.5],
    )
    add_paragraph(doc, "【モジュール定数】", bold=True)
    add_bullets(doc, [
        'ALLOWED_EXTENSIONS = ["jpg", "jpeg", "png"]',
        "MAX_FILE_SIZE = 5MB（5 * 1024 * 1024）",
        "JPEG_QUALITY = 90",
    ])

    add_h3(doc, "8.4.2 json_normalizer.py")
    add_table(doc,
        headers=["関数名", "性質", "用途"],
        rows=[
            ["normalize_to_contact_dict", "純関数", "raw_json → Contact 用辞書 + confidence_map"],
            ["calc_orientation_adjusted_confidence_map", "純関数", "orientation 別に confidence を補正（v1.3.0 新規）"],
            ["_extract_value_and_confidence", "純関数（内部）", "フィールド構造から値と信頼度を抽出"],
        ],
        col_widths_cm=[6.0, 4.0, 6.5],
    )
    add_paragraph(doc, "【防御的実装の方針】", bold=True)
    add_bullets(doc, [
        "構造想定外でも例外を投げず、ベストエフォートで処理",
        "例外は card_index が cards 配列の範囲外の場合のみ ValueError を送出",
    ])

    add_h3(doc, "8.4.3 has_minimum_info.py")
    add_table(doc,
        headers=["関数名", "性質", "用途"],
        rows=[
            ["has_minimum_info", "純関数", "Contact として最低限の情報があるかを判定"],
            ["_stripped", "純関数（内部）", "None / 空白 / \\n を空文字扱いに正規化"],
        ],
        col_widths_cm=[4.5, 4.5, 7.5],
    )
    add_paragraph(doc, "【判定ロジック】", bold=True)
    add_bullets(doc, [
        '_REQUIRED_FIELD = "full_name"（必須）',
        '_AT_LEAST_ONE_OF = ("company", "email", "phone", "mobile")（少なくとも 1 つ非空）',
        "両条件を満たす場合のみ True を返す",
        "バリデーションではない（メールアドレスの形式妥当性などは責務外）",
    ])

    # 8.5 Task 層
    add_h2(doc, "8.5 Task 層の設計（cards/tasks/）", mark_revision=True)
    add_paragraph(doc,
        "Task 層は副作用あり処理（DB 書き込み・API 呼び出し・ファイル書き込み）を担当する。"
        "Service 層との違いは「外部世界との通信を伴うか」"
    )

    add_h3(doc, "8.5.1 PipelineCoordinator", mark_revision=True)
    add_paragraph(doc,
        "役割：1 枚の OriginalImage に対して、OCR パイプライン全体を統括する。"
        "実装形式はクラスとし、self.original_image / self.error_messages / self._cards_by_index 等の状態を保持する。"
    )
    add_paragraph(doc, "【主要メソッド】", bold=True)
    add_table(doc,
        headers=["メソッド名", "性質", "用途"],
        rows=[
            ["__init__", "初期化", "OriginalImage を受け取り、状態を初期化"],
            ["run_pipeline", "副作用あり", "パイプライン本体（OCR・DB・ファイル）"],
            ["_process_card", "副作用あり（内部）", "1 枚の card 単位で処理（OCR・BC・Contact 作成）"],
            ["_save_dev_image", "副作用あり（DEBUG 専用）", "開発時のみ中間画像を保存"],
            ["_load_combined_schema / _load_card_item_schema", "副作用あり（キャッシュ）", "JSON Schema 読み込み"],
        ],
        col_widths_cm=[6.0, 4.0, 6.5],
    )

    add_h4(doc, "8.5.1.1 状態遷移", mark_revision=True)
    add_paragraph(doc,
        "v1.3.0 で CAS 方式に変更したため、run_pipeline 開始時には既に status=processing が確定している。"
        "正常・異常両パスで必ず以下のいずれかに遷移すること。"
    )
    # ★修正④：extracted の条件
    add_table(doc,
        headers=["遷移後 status", "条件"],
        rows=[
            ["extracted", "名刺検出 ≧1 件。has_minimum_info で全 card が弾かれた場合も BusinessCard 0 件の状態で extracted に遷移する（v1.3.3 改訂）"],
            ["garbage", "名刺検出 0 件（detect_cards 結果が空）"],
            ["failed", "OCR API 失敗・スキーマ検証失敗・予期しない例外発生"],
        ],
        col_widths_cm=[4.0, 12.5],
        revision_row_indices={0},
    )
    add_paragraph(doc,
        "【pending に戻る経路の例外】pending に戻る経路は run_pipeline 内には存在しない。"
        "ただし stuck sweeper（8.5.4.2）は例外的に processing → pending に差し戻す。"
        "差し戻し時には _cleanup_stuck_record で再処理可能な状態にクリーンアップする"
        "（既存 BusinessCard・Contact・孤立 Person・card_image を削除）ため、再処理時の整合性は保たれる。"
    )

    add_h4(doc, "8.5.1.2 run_pipeline 冒頭の防御チェック（v1.3.1 改訂）", mark_revision=True)
    add_paragraph(doc,
        "v1.3.0 では status が processing 以外で run_pipeline が呼ばれた場合、警告ログを出して処理続行していた。"
        "v1.3.1 では CAS 前提の安全策として、status≠processing なら failed に落として早期 return する。"
        "理由：CAS が成立している前提で動作するロジックなので、前提が崩れた状態で続行すると"
        "二重処理・データ破損のリスクがある。"
    )

    add_h4(doc, "8.5.1.3 _process_card の例外捕捉範囲", mark_revision=True)
    add_paragraph(doc,
        "v1.2.x では except OcrApiError のみで捕捉していたため、JSONDecodeError や KeyError などの想定外例外が漏れ、"
        "1 枚の card 失敗が OriginalImage 全体の STATUS_FAILED に波及していた。"
        "v1.3.0 で except Exception で広く捕捉する形に修正し、当該 card のみ失敗扱いとして他 card への波及を防ぐ。"
    )

    add_h4(doc, "8.5.1.4 ファイル保存方式（DB先・ファイル後）", mark_revision=True)
    add_paragraph(doc,
        "v1.2.x ではファイル書き込み → DB 書き込みの順序だったため、両者の間でプロセスが SIGKILL された場合に"
        "「ファイルあり・DB 行なし」の孤児ファイルが残るリスクがあった。"
        "v1.3.0 で DB 先行・ファイル後方式に変更した。"
    )
    add_image(doc, "E_file.png", "図 8-1：ファイル保存フロー（DB 先・ファイル後）")
    add_paragraph(doc, "【処理フロー概要】", bold=True)
    add_bullets(doc, [
        "①トランザクション内で BusinessCard を card_image=None で先行作成",
        "②同じトランザクション内で Contact / Person 作成",
        "③一時ファイル（.jpg.tmp）に画像を書き込み",
        "④transaction.on_commit() でリネーム＆UPDATE 処理を登録",
        "⑤コミット成功時に on_commit が実行され、.tmp → .jpg にリネーム",
        "⑥card_image フィールドを UPDATE してパスを記録",
    ])
    add_paragraph(doc, "【失敗パターン別の挙動】", bold=True)
    add_table(doc,
        headers=["失敗タイミング", "挙動"],
        rows=[
            ["atomic 内で例外", "os.unlink(tmp) で一時ファイルを削除し、DB 行はロールバック"],
            ["atomic コミット後・rename 失敗", "BusinessCard は残るが card_image=None のまま。reconcile コマンドで欠損検出可能"],
            ["rename 成功・UPDATE 失敗", "renamed フラグにより final_path を削除し、card_image=None を保つ"],
        ],
        col_widths_cm=[5.5, 11.0],
    )
    add_paragraph(doc,
        "【テスト時の注意】transaction.on_commit() は Django TestCase 内では呼ばれない"
        "（テスト全体がロールバックされるため）。テストする際は TransactionTestCase を使うか、"
        "TestCase.captureOnCommitCallbacks()（Django 4.1+）を使用すること。"
    )

    add_h4(doc, "8.5.1.5 raw_json の部分保存", mark_revision=True)
    add_paragraph(doc,
        "v1.2.x では failed パスで raw_json を設定していなかったため、一部 card が成功した状態で後段例外が発生すると、"
        "「BusinessCard が紐付いているのに raw_json=None」となり v1.2.0 の不変条件が壊れていた。"
        "v1.3.0 で except ブランチで _cards_by_index から部分 raw_json を組み立てて保存するよう修正。"
    )
    add_bullets(doc, [
        "_cards_by_index にエントリがある場合：部分 raw_json を組み立てて保存",
        "_cards_by_index が空（detect_cards 自体が失敗）：raw_json=None のまま保持",
        "これにより不変条件「BusinessCard.card_index → raw_json[\"cards\"][i]」を保全",
    ])

    add_h4(doc, "8.5.1.6 final save の update_fields 指定", mark_revision=True)
    add_paragraph(doc,
        "v1.2.x では original_image.save() を update_fields 指定なしで呼んでいたため、"
        "全列 UPDATE による他プロセスとの競合書き込みリスクがあった。"
        "v1.3.0 で明示的に update_fields を指定し、atomic でラップする形に修正。"
        "更新フィールド：status / raw_json / error_message / detected_count / updated_at"
    )

    # 8.5.2 OcrService
    add_h3(doc, "8.5.2 OcrService", mark_revision=True)
    add_paragraph(doc,
        "役割：OCR バックエンドとの通信を抽象化。標準 JSON を返す。"
        "v1.3.x では Claude Sonnet 4.6 のみだが、v2.0.0 で他バックエンド追加時の修正範囲を限定するため"
        "クラスとして実装する。"
    )
    add_paragraph(doc, "【主要メソッド】", bold=True)
    add_table(doc,
        headers=["メソッド名", "性質", "用途"],
        rows=[
            ["__init__(api_key, model)", "初期化", "クライアントは遅延生成"],
            ["run_ocr", "副作用あり（API 呼出）", "画像 → 標準 JSON。失敗時 OcrApiError を raise"],
            ["_prepare_image", "純関数", "画像リサイズ・base64 変換"],
            ["_load_schema / _load_tool_input_schema / _load_prompt", "副作用あり（キャッシュ）", "ファイル読み込み（初回のみ）"],
            ["_extract_tool_use_input / _extract_api_response", "純関数", "レスポンスから必要データを抽出"],
        ],
        col_widths_cm=[6.5, 4.0, 6.0],
    )
    add_paragraph(doc, "【クラス定数】", bold=True)
    add_bullets(doc, [
        'DEFAULT_MODEL = "claude-sonnet-4-6"',
        "DEFAULT_MAX_TOKENS = 4096",
        "MAX_LONG_EDGE = 1568（画像長辺の最大ピクセル）",
        'SCHEMA_VERSION = "1.3.0"',
        'TOOL_NAME = "extract_business_card"',
    ])
    add_h4(doc, "8.5.2.1 API タイムアウト設定（v1.3.0 新規）", mark_revision=True)
    add_paragraph(doc,
        "v1.2.x では Anthropic SDK のデフォルトタイムアウト（数百秒）に依存していた。"
        "cron が 1〜5 分間隔で起動される設計のため、タイムアウトした worker がプロセスを長時間占有すると、"
        "後続 cron と重なって全 worker がブロックされるリスクがあった。"
        "v1.3.0 でクライアント生成時に明示的に 60 秒を設定。"
    )
    add_paragraph(doc, "【関連例外】", bold=True)
    add_bullets(doc, [
        "OcrApiError(Exception)：API 失敗を表すカスタム例外。run_pipeline 側で except Exception により広く捕捉される",
    ])

    # 8.5.3 CardCropper ★修正⑤⑥
    add_h3(doc, "8.5.3 CardCropper", mark_revision=True)
    add_paragraph(doc,
        "役割：透視変換済みの warped_image（PIL Image）を受け取り、"
        "名刺画像として MEDIA_ROOT 配下に JPEG として保存する。"
        "bbox → warped_image の変換は opencv_detector 側が担う。"
        "実装形式は関数（クラスではない）。状態を持たず、入力 → 出力の純粋な変換処理のため。"
        "（v1.3.3 改訂）"
    )
    add_paragraph(doc, "【主要関数】", bold=True)
    add_table(doc,
        headers=["関数名", "性質", "用途"],
        rows=[
            ["save_card_image", "副作用あり", "名刺画像を直接保存（旧フロー用）"],
            ["save_card_image_tmp", "副作用あり（一時）", "一時ファイル（.jpg.tmp）に保存。on_commit でリネーム前提（v1.3.0 新規）"],
            ["_build_save_path", "副作用あり（内部）", "保存パスを生成（時刻ベース）"],
        ],
        col_widths_cm=[5.0, 4.5, 7.0],
        revision_row_indices={1},
    )
    add_paragraph(doc, "【最低画像サイズ基準】", bold=True)
    add_paragraph(doc,
        "§6.2 を参照。判定は opencv_detector._warp_card() 側で行い、"
        "width < 100 または height < 50 の場合は warped_image を生成しない（CardCropper には到達しない）。"
        "（v1.3.3 改訂：v1.3.2 までは「CardCropper が (False, None, '画像サイズが最低基準未満') を返す」と"
        "誤って記述していた）"
    )

    # 8.5.4 process_pending
    add_h3(doc, "8.5.4 process_pending 管理コマンド", mark_revision=True)
    add_paragraph(doc,
        "cron で 1〜5 分間隔で起動される。v1.2.2 で楽観ロック方式を方針として記述していたが、"
        "実装上の CAS（Compare-And-Swap）が成立していなかったため、v1.3.0 で根本修正。"
        "v1.3.1 で stuck sweeper の再処理クリーンアップを追加した。"
    )
    add_paragraph(doc, "worker の正常時・異常時の挙動を以下の図に示す。")
    add_image(doc, "D_timeline.png", "図 8-2：worker のタイムライン（正常時・異常終了時・stuck sweeper 発動）")

    add_h4(doc, "8.5.4.1 _claim_lock の CAS 化", mark_revision=True)
    add_paragraph(doc,
        "【v1.2.2 の問題】filter(status='pending').update(updated_at=now) では status を変えていないため、"
        "Worker A と Worker B が同じレコードに対して両方とも updated == 1 を得てしまい、"
        "CAS が成立せず多重処理が発生していた。"
    )
    add_paragraph(doc,
        "【v1.3.0 の解決】_claim_lock 内の UPDATE で status=PROCESSING と claimed_at=now を同時に書き込むことで"
        "CAS を成立させる。Worker A が UPDATE 成功した瞬間に status=processing に変わるため、"
        "Worker B の UPDATE は条件不一致で 0 件となる。"
    )

    add_h4(doc, "8.5.4.2 stuck sweeper（v1.3.1 改訂）", mark_revision=True)
    add_paragraph(doc,
        "CAS で processing に遷移したレコードは、worker が SIGKILL や OOM で異常終了した場合に"
        "「誰も拾わない孤児レコード」として永久に残るリスクがある。"
        "v1.3.0 で handle() 先頭に stuck sweeper を追加し、v1.3.1 で再処理クリーンアップを追加した。"
    )
    add_paragraph(doc, "【しきい値】", bold=True)
    add_bullets(doc, [
        "settings.OCR_STUCK_THRESHOLD_MINUTES = 30（デフォルト 30 分）",
        "Claude API のタイムアウト 60 秒・処理時間の余裕を考慮した値",
        "運用状況に応じて settings.py で調整可能",
    ])
    add_paragraph(doc, "【動作（v1.3.1 改訂）】", bold=True)
    add_paragraph(doc,
        "stuck レコードを単純に pending に戻すだけでは、既存の BusinessCard が残ったまま再処理されて"
        "UniqueConstraint(original_image, card_index) 違反になる。"
        "v1.3.1 では _cleanup_stuck_record で「再処理可能な状態」までクリーンアップしてから pending に戻す。"
    )

    add_h4(doc, "8.5.4.3 _cleanup_stuck_record（v1.3.1 新規）", mark_revision=True)
    add_paragraph(doc,
        "stuck レコードに対する再処理クリーンアップ関数。"
        "以下の処理を transaction.atomic() 内で実行する。"
    )
    add_bullets(doc, [
        "BusinessCard に紐付く Contact から person_id を集合に記録（後で孤立判定）",
        "card_image ファイルを os.unlink で削除（FileNotFoundError は握りつぶす）",
        "BusinessCard を delete()（CASCADE で Contact / ContactFieldConfidence も削除）",
        "記録した person_id について、他 Contact から参照されていなければ Person を delete()",
        "OriginalImage を raw_json=None / error_message=\"\" / status=PENDING / claimed_at=None / detected_count=0 でリセット",
    ])
    add_paragraph(doc,
        "【動作上の注意】ファイル削除は transaction.atomic() 内で実行している。"
        "DB ロールバック時にファイルだけが消えて DB 行が残るパターンが起きうるが、"
        "stuck レコードはそもそも異常状態であり、reconcile_card_images で後から欠損検出できるため、許容範囲とする。"
    )

    add_h4(doc, "8.5.4.4 処理上限")
    add_bullets(doc, [
        "オプション --limit で最大処理件数を指定（デフォルト 10）",
        "SQLite 開発中は 1 件ずつでも可",
        "本番 DB 移行時に select_for_update(skip_locked=True) を検討",
    ])

    # 8.5.5 retry_failed_ocr
    add_h3(doc, "8.5.5 retry_failed_ocr 管理コマンド", mark_revision=True)
    add_paragraph(doc, "開発・運用ツール。エンドユーザー向けではない。")
    add_paragraph(doc, "【オプション】", bold=True)
    add_bullets(doc, [
        "--all：すべての failed を対象",
        "--id <UUID>：特定の OriginalImage を対象（--all と排他）",
        "--limit：最大件数",
        "--dry-run：実行せず対象を表示するだけ",
    ])
    add_paragraph(doc, "【対象範囲】", bold=True)
    add_bullets(doc, [
        "status=failed AND BusinessCard 0 件 の OriginalImage",
        "部分成功（status=extracted で error_message あり）は対象外（v2.0.0 で検討）",
        "garbage は対象外（OCR 結果として正しい判定のため）",
    ])
    add_paragraph(doc, "【処理内容】", bold=True)
    add_bullets(doc, [
        "対象 OriginalImage の raw_json を null クリア",
        "error_message をクリア・detected_count を 0 にリセット",
        "status を pending に戻す（次の cron 起動で process_pending が拾う）",
    ])
    add_h4(doc, "8.5.5.1 race 縮小（v1.3.0 改訂）", mark_revision=True)
    add_paragraph(doc,
        "v1.2.x では SELECT で対象を取得した後、UPDATE で id__in=target_ids のみ指定していたため、"
        "SELECT と UPDATE の間に他 worker（process_pending 等）が status を変える可能性があった。"
        "v1.3.0 で UPDATE 側の filter にも status=STATUS_FAILED 条件を追加して race を縮小。"
    )

    # 8.5.6 jsonschema 再検証
    add_h3(doc, "8.5.6 jsonschema 再検証")
    add_paragraph(doc,
        "Tool Use で形式強制していても、保存前に jsonschema で再検証する。"
        "理由：将来のバックエンド追加時（GPT-4o / Vision 等）も同じ検証を通すため、"
        "また Claude の挙動変化への防御策。"
    )
    add_paragraph(doc, "【検証失敗時の挙動（現状）】", bold=True)
    add_bullets(doc, [
        "card 単位スキーマ違反：当該 card のみスキップ（BusinessCard / Contact 作成なし、他 card は処理続行）",
        "combined（全体）スキーマ違反：v1.3.x 時点では warning ログのみで保存続行（v1.4.0 で failed 化予定、§12.3）",
    ])

    # 8.5.7 reconcile_card_images
    add_h3(doc, "8.5.7 reconcile_card_images 管理コマンド", mark_revision=True)
    add_paragraph(doc,
        "DB と MEDIA_ROOT の整合性を検査する管理コマンド。"
        "8.5.1.4 のファイル保存方式変更（DB 先・ファイル後）により孤児ファイルの発生可能性は理論上低くなったが、"
        "ディスク障害・手動削除・バックアップ復元時のズレ・stuck cleanup 時のロールバック等による不整合は"
        "依然として起こり得る。"
    )
    add_paragraph(doc, "検出ロジックを以下に示す。Set 演算による O(n) の効率的な処理。")
    add_image(doc, "F_reconcile.png", "図 8-3：reconcile_card_images の検出ロジック")

    add_h4(doc, "8.5.7.1 オプション")
    add_bullets(doc, [
        "（既定）dry-run：検出・レポートのみ、ファイル操作なし",
        "--apply：実修復を実行（孤児を MEDIA_ROOT/cards/_orphan/ に退避）",
    ])
    add_h4(doc, "8.5.7.2 検出対象")
    add_table(doc,
        headers=["検出種別", "条件"],
        rows=[
            ["孤児ファイル", "MEDIA_ROOT 配下にファイルが存在するが、BusinessCard.card_image にパスが記録されていない"],
            ["欠損ファイル", "BusinessCard.card_image にパスが記録されているが、実ファイルが存在しない"],
        ],
        col_widths_cm=[3.5, 13.0],
    )
    add_h4(doc, "8.5.7.3 動作モード")
    add_table(doc,
        headers=["モード", "孤児ファイル", "欠損ファイル"],
        rows=[
            ["dry-run（既定）", "stdout に絶対パスを出力", "stdout に BusinessCard ID + パスを出力"],
            ["--apply", "MEDIA_ROOT/cards/_orphan/ に os.rename で移動。同名衝突時はマイクロ秒精度のタイムスタンプ付与", "ログ出力のみ（自動修復なし）"],
        ],
        col_widths_cm=[3.5, 8.0, 5.0],
    )
    add_h4(doc, "8.5.7.4 検出ロジック")
    add_bullets(doc, [
        "DB スキャン：BusinessCard の card_image 相対パスを Set 化",
        "FS スキャン：os.walk で MEDIA_ROOT/cards/ を再帰走査",
        "FS スキャン時：_orphan/ サブディレクトリは枝刈り、.tmp ファイルは除外",
        "孤児 = fs_paths - db_paths",
        "欠損 = db_paths - fs_paths",
    ])
    add_h4(doc, "8.5.7.5 運用上の注意")
    add_paragraph(doc,
        "【パイプライン非稼働時間帯の実行を推奨】--apply 実行中にパイプラインが _finalize_card_image"
        "（on_commit）を実行していた場合、.tmp → .jpg のリネーム直後に reconcile が .jpg を「孤児」と"
        "誤検出して退避する可能性がある。cron が走っていない夜間などに実行することを推奨する。"
    )

    # 8.6 Management Commands 一覧
    add_h2(doc, "8.6 Management Commands 一覧", mark_revision=True)
    add_table(doc,
        headers=["コマンド名", "v1.3.x 状況", "用途"],
        rows=[
            ["process_pending", "v1.3.1 改訂", "cron 起動。pending 画像を OCR 処理。CAS + stuck sweeper"],
            ["retry_failed_ocr", "v1.3.0 改訂", "failed を pending に戻す。開発・運用ツール"],
            ["reconcile_card_images", "v1.3.0 新規", "DB↔MEDIA_ROOT 整合検査"],
            ["dev_for_reset_ocr", "既存", "開発用：BusinessCard 削除＋OriginalImage を pending リセット"],
        ],
        col_widths_cm=[5.0, 3.5, 8.0],
        revision_row_indices={0, 1, 2},
    )


def section_chapter9(doc) -> None:
    add_h1(doc, "第9章 関数命名規則と性質明記")
    add_h2(doc, "9.1 関数の3分類")
    add_bullets(doc, [
        "純関数：DB を一切触らない、副作用なし、同じ入力で同じ出力",
        "準関数：DB を読むが書かない、外部世界に副作用なし",
        "副作用あり関数：DB 書き込み・例外送出・API 呼び出し・ファイル書き込みなど",
    ])
    add_h2(doc, "9.2 プレフィックス")
    add_table(doc,
        headers=["プレフィックス", "性質", "例"],
        rows=[
            ["normalize_* / to_* / calc_* / is_* / has_*", "純関数", "normalize_to_contact_dict / has_minimum_info"],
            ["find_* / get_* / search_*", "準関数", "find_pending_images"],
            ["validate_*", "副作用あり（例外）", "validate_image"],
            ["convert_* / save_* / create_* / update_* / delete_*", "副作用あり（変換・DB 書込）", "convert_to_jpeg / save_card_image"],
            ["run_* / process_* / send_*", "副作用あり（複合処理）", "run_ocr / run_pipeline / process_pending"],
            ["retry_*", "副作用あり（再投入）", "retry_failed_ocr"],
        ],
        col_widths_cm=[6.5, 4.5, 5.5],
    )
    add_h2(doc, "9.3 docstring 性質明記")
    add_paragraph(doc, "関数の docstring 冒頭に性質を明記する：")
    add_bullets(doc, [
        "[性質] 純関数 / 準関数 / 副作用あり",
        "[入力] パラメータの説明",
        "[出力] 戻り値の説明",
        "[例外] 送出する例外の種類と条件",
    ])


def section_chapter10(doc) -> None:
    add_h1(doc, "第10章 エラーハンドリング・cards 処理ルール")
    # ★修正④：C3 行
    add_h2(doc, "10.1 cards 処理ルール一覧")
    add_table(doc,
        headers=["ID", "ケース", "処理"],
        rows=[
            ["C1", "cards 配列が空", "status=garbage、BusinessCard 作成なし"],
            ["C2", "全 card が is_business_card=false", "status=garbage、BusinessCard 作成なし"],
            ["C3", "全 card が has_minimum_info で弾かれた",
                "status=extracted（BusinessCard 0 件）、OriginalImage.error_message に「全 card has_minimum_info 不満」を記録（v1.3.3 改訂）"],
            ["C4", "confidence=low/medium のフィールドあり", "ContactFieldConfidence に記録（high は記録しない）"],
            ["C5", "切り抜き失敗（card_image=null）",
                "has_minimum_info を満たせば BusinessCard 作成。理由は OriginalImage.error_message に記録"],
            ["C6", "OCR 例外発生（OcrApiError 以外含む）",
                "当該 card のみ失敗扱い、他 card への波及なし（v1.3.0 改訂）"],
        ],
        col_widths_cm=[1.2, 5.5, 9.8],
        revision_row_indices={2},
    )
    add_paragraph(doc,
        "※ 現在の実装では、has_minimum_info で全 card が弾かれた場合も status=extracted（BusinessCard 0 件）"
        "となるため、UI 上「完了なのに名刺データが表示されない」状態が生じる。"
        "v1.4.0 で garbage との境界定義を見直し予定（§12.3）。",
        italic=True,
    )

    add_h2(doc, "10.2 部分失敗時の status / error_message 仕様")
    add_bullets(doc, [
        "少なくとも 1 件の BusinessCard 作成成功 → status=extracted、error_message に失敗 card の理由を集約",
        "全 card 失敗 → status=failed、error_message に詳細を記録",
        "error_message は TextField なので複数行記録可能（card_index ごとに改行）",
    ])


def section_chapter11(doc) -> None:
    add_h1(doc, "第11章 非機能要件", mark_revision=True)
    add_h2(doc, "11.1 性能要件")
    add_bullets(doc, [
        "画像アップロード → 名刺一覧表示まで：画像 1 枚あたり 30 秒以内",
        "OCR API 1 回あたりのタイムアウト：60 秒（★v1.3.0 設定）",
        "process_pending の cron 起動間隔：1〜5 分",
        "stuck sweeper のしきい値：30 分（★v1.3.0 新規）",
    ])
    add_h2(doc, "11.2 信頼性要件")
    add_bullets(doc, [
        "多重起動対策：CAS 楽観ロック方式（★v1.3.0 改訂）",
        "異常終了対策：stuck sweeper による再処理クリーンアップ（★v1.3.1 改訂）",
        "整合性検査：reconcile_card_images コマンドによる定期監視（★v1.3.0 新規）",
    ])
    add_h2(doc, "11.3 設定値一覧（v1.3.0 新規）", mark_revision=True)
    add_table(doc,
        headers=["設定キー", "既定値", "用途"],
        rows=[
            ["OCR_STUCK_THRESHOLD_MINUTES", "30", "stuck sweeper のしきい値（分）（v1.3.0 新規）"],
            ["ANTHROPIC_API_KEY", "(必須)", "Claude API の API キー"],
            ["MEDIA_ROOT", "(必須)", "画像ファイル保存先"],
        ],
        col_widths_cm=[6.0, 3.0, 7.5],
        revision_row_indices={0},
    )
    # ★修正②：認証
    add_h2(doc, "11.4 セキュリティ要件")
    add_bullets(doc, [
        "認証（今後の導入予定）：Django 標準認証 + LoginRequiredMixin。v1.3.x では get_current_user() 仮実装で認証スキップ中、Phase 4 で本格導入（v1.3.3 改訂）",
        "CSRF：Django 標準保護",
        "画像アップロードはバリデーション必須（image_processor.validate_image を UploadForm.clean_image() 経由で呼ぶ）",
    ])


def section_chapter12(doc) -> None:
    add_h1(doc, "第12章 制約事項・将来の拡張", mark_revision=True)
    add_h2(doc, "12.1 v1.3.4 の制約事項")
    add_bullets(doc, [
        "OCR 処理の起動ユーザーアクションは画像アップロードのみ",
        "失敗した OriginalImage の retry 機能はユーザーに提供しない",
        "retry_failed_ocr は開発・運用ツールとして位置付け（本番ではエンドユーザーに見せない）",
        "OriginalImage.status に processing 状態を追加（CAS 中間状態）",
        "OriginalImage.claimed_at は CAS で processing 遷移時のみ書き込まれ、stuck sweeper cleanup 時に None に戻る",
        "ファイル保存方式は DB 先・ファイル後（on_commit でリネーム）。テスト時は TransactionTestCase か captureOnCommitCallbacks() が必要",
        "Claude API タイムアウトは 60 秒固定（v1.3.x 時点では設定値化していない）",
        "reconcile_card_images の --apply はパイプライン非稼働時間帯に実行すること",
        "画像アップロード上限：5MB / JPEG・PNG のみ（HEIC は v1.3.x では受付不可）",
        "認証は仮実装状態（cards.views.get_current_user() がデフォルトユーザーを返す）。Phase 4 で LoginRequiredMixin による本格導入予定（v1.3.3 改訂）",
        "garbage と extracted の境界が現状不揃い（has_minimum_info 全落ちが extracted 扱い）。v1.4.0 で見直し予定（v1.3.3 改訂）",
    ])

    add_h2(doc, "12.2 stuck sweeper cleanup の副作用", mark_revision=True)
    add_paragraph(doc,
        "_cleanup_stuck_record によって以下の副作用が発生する。いずれも軽微な影響であり、"
        "reconcile_card_images で後から検出可能なため許容範囲とする。"
    )
    add_paragraph(doc, "【副作用1：DB ロールバック時のファイル消失】", bold=True)
    add_bullets(doc, [
        "ファイル削除を transaction.atomic() 内で実行している",
        "DB トランザクションがロールバックされた場合、削除済みファイルは復元されない",
        "結果：DB 行に card_image パスが残り、実ファイルが存在しない欠損状態になりうる",
        "対処：reconcile_card_images で「欠損ファイル」として検出可能",
    ])
    add_paragraph(doc, "【副作用2：孤立 Person の発生は v1.3.1 で解消】", bold=True)
    add_bullets(doc, [
        "v1.3.0 までは BusinessCard.delete() による Contact CASCADE 削除後、Person が孤立していた",
        "v1.3.1 で _cleanup_stuck_record 内で「他 Contact から参照されていない Person」を検出して削除するロジックを追加",
    ])

    add_h2(doc, "12.3 v1.4.0 以降で予定する変更")
    add_bullets(doc, [
        "認証の本格導入（LoginRequiredMixin / 適切なユーザーモデル）",
        "garbage と extracted の境界定義見直し（has_minimum_info 全落ち時の扱い）",
        "HEIC 受付対応・最大ファイルサイズの引き上げ",
        "Claude API タイムアウトを settings.py の設定値化",
        "reconcile_card_images の自動実行（cron / Celery beat 経由）と進行中レコード除外ガード",
        "schema 検証失敗時の挙動を warning 止まりから明示的 failed 化に変更",
        "retry_failed_ocr の対象を部分成功（status=extracted で error_message あり）まで拡張",
        "retry_count フィールド追加：N 回失敗で自動的に failed 確定（無限ループ対策）",
    ])

    add_h2(doc, "12.4 v2.0.0 で予定する拡張")
    add_bullets(doc, [
        "OCR バックエンドの切替対応（GPT-4o / Vision / EasyOCR）",
        "Person 統合・自動マージ機能",
        "OriginalImage / BusinessCard の削除機能（UI 提供）",
        "再 OCR 設計（raw_json 上書きなし）",
        "SNS 別テーブル化",
        "メール配信・案件管理・スケジュール管理・設備予約への発展",
    ])


def section_appendix(doc) -> None:
    add_h1(doc, "巻末別表")
    add_h2(doc, "別表A：モデル名・フィールド名 日本語/英語対照表")
    add_h3(doc, "A.1 モデル名")
    add_table(doc,
        headers=["日本語表記", "コード上の名前", "補足"],
        rows=[
            ["元画像DB", "OriginalImage", "ユーザーアップロード画像"],
            ["名刺DB", "BusinessCard", "切り抜き済み名刺画像"],
            ["コンタクトDB", "Contact", "名刺ごとのスナップショット"],
            ["人物DB", "Person", "人物本体"],
            ["信頼度メタDB", "ContactFieldConfidence", "Contact フィールド信頼度"],
        ],
        col_widths_cm=[4.0, 6.0, 6.5],
    )
    add_h3(doc, "A.2 主要フィールド名")
    add_table(doc,
        headers=["日本語表記（完全修飾名）", "コード上の名前", "v1.3.x 状況"],
        rows=[
            ["元画像DB．処理ステータス", "OriginalImage.status", "processing 追加"],
            ["元画像DB．取得時刻", "OriginalImage.claimed_at", "v1.3.0 新規"],
            ["元画像DB．OCR結果JSON", "OriginalImage.raw_json", "既存"],
            ["元画像DB．検出件数", "OriginalImage.detected_count", "既存"],
            ["名刺DB．カード画像", "BusinessCard.card_image", "nullable"],
            ["名刺DB．配列インデックス", "BusinessCard.card_index", "既存"],
            ["名刺DB．向き", "BusinessCard.orientation", "v1.3.0 新規（5 値）"],
            ["コンタクトDB．名前", "Contact.full_name", "既存"],
            ["信頼度メタDB．フィールド名", "ContactFieldConfidence.field_name", "Contact 側名"],
        ],
        col_widths_cm=[6.5, 6.0, 4.0],
        revision_row_indices={0, 1, 6},
    )

    add_h2(doc, "別表B：管理コマンド一覧（v1.3.4 完全版）")
    add_table(doc,
        headers=["コマンド", "オプション", "用途"],
        rows=[
            ["process_pending", "--limit", "cron 起動。stuck sweeper + CAS + run_pipeline"],
            ["retry_failed_ocr", "--all / --id / --limit / --dry-run", "failed を pending に戻す"],
            ["reconcile_card_images", "--apply", "DB↔MEDIA_ROOT 整合検査・修復"],
            ["dev_for_reset_ocr", "--all / --id / --limit / --dry-run", "開発用リセット"],
        ],
        col_widths_cm=[5.0, 5.5, 6.0],
    )

    add_h2(doc, "別表C：Phase 3-2 実装スコープ（v1.3.4 改訂）", mark_revision=True)
    add_paragraph(doc, "Phase 3-2 で実装したファイル一覧と v1.3.x における変更状況。")
    add_table(doc,
        headers=["#", "ファイル", "保存先", "状況", "変更内容"],
        rows=[
            ["1", "views.py", "cards/", "v1.3.3 修正",
                "UploadView リダイレクト先（originals:original_detail）と認証仮実装の明記"],
            ["2", "pipeline_coordinator.py", "cards/tasks/", "v1.3.1 改訂",
                "防御チェック・状態遷移整合・raw_json 部分保存・update_fields 指定・DB 先ファイル後・例外捕捉拡大"],
            ["3", "ocr_service.py", "cards/tasks/", "v1.3.0 改訂", "API タイムアウト 60 秒設定"],
            ["4", "card_cropper.py", "cards/tasks/", "v1.3.0 改訂", "save_card_image_tmp 追加"],
            ["5", "process_pending.py", "cards/management/commands/", "v1.3.1 改訂",
                "CAS 化・stuck sweeper・_cleanup_stuck_record 追加"],
            ["6", "retry_failed_ocr.py", "cards/management/commands/", "v1.3.0 改訂",
                "UPDATE filter に status 追加・detected_count リセット"],
            ["7", "reconcile_card_images.py", "cards/management/commands/", "v1.3.0 新規",
                "DB↔MEDIA_ROOT 整合検査"],
            ["8", "json_normalizer.py", "cards/services/", "v1.3.0 改訂",
                "calc_orientation_adjusted_confidence_map 追加"],
            ["9", "has_minimum_info.py", "cards/services/", "既存", "v1.2.x から変更なし"],
            ["10", "image_processor.py", "cards/services/", "既存", "validate_image / convert_to_jpeg"],
            ["11", "opencv_detector.py", "cards/services/detectors/", "v1.3.3 明記",
                "最低画像サイズ判定（width<100 or height<50）の所在をここと仕様書で明記"],
        ],
        col_widths_cm=[1.0, 4.5, 5.0, 2.5, 4.5],
        revision_row_indices={0, 1, 2, 3, 4, 5, 6, 7, 10},
    )


def section_history(doc) -> None:
    add_h1(doc, "改訂履歴")
    add_table(doc,
        headers=["バージョン", "日付", "改訂内容", "改訂者"],
        rows=[
            ["v1.0.0", "2026/04/27", "新規作成", "たんたん"],
            ["v1.0.1", "2026/04/27", "GPT レビュー指摘 4 点を反映", "たんたん"],
            ["v1.1.0", "2026/04/28", "Contact 配列廃止・ContactFieldConfidence 追加・services 統合", "たんたん"],
            ["v1.2.0", "2026/04/28", "raw_json を OriginalImage に集約・BusinessCard 簡素化・card_index 追加", "たんたん"],
            ["v1.2.1", "2026/04/28",
                "cards 処理ルール明文化・OcrBackend 抽象化方針・JSON Schema 管理・card_image null 許容", "たんたん"],
            ["v1.2.2", "2026/04/28", "Phase 3-2 着手前確定版（実装方針・ディレクトリ構成・命名規則・運用方針）", "たんたん"],
            ["v1.3.0", "2026/04/30",
                "OCR パイプライン堅牢性修正（CAS 化・stuck sweeper・DB先ファイル後・reconcile コマンド・"
                "API タイムアウト・例外捕捉拡大・raw_json 部分保存）", "たんたん"],
            ["v1.3.1", "2026/04/30",
                "完結版・全文掲載。run_pipeline 防御チェック・stuck cleanup（既存 BusinessCard / Contact / 孤立 Person 削除）", "たんたん"],
            ["v1.3.2", "2026/04/30",
                "コードスニペット削除・概念図 6 枚追加・View/Service/Task 命名整合・実装との差異解消（5MB / NOT NULL / source_value 削除等）", "たんたん"],
            ["v1.3.3", "2026/05/01",
                "v1.3.2 レビュー指摘 6 件修正（UploadView リダイレクト・認証の仮実装状況・"
                "JSON Schema ファイル名・garbage 定義統一・最低サイズ判定と CardCropper 役割説明）", "たんたん"],
            ["v1.3.4", "2026/05/01",
                "v1.3.3 見栄え 3 点修正（Heading スタイル、表ヘッダー薄青背景、★改訂セル黄色背景）。"
                "コンテンツ変更なし", "たんたん"],
        ],
        col_widths_cm=[2.5, 2.5, 9.5, 2.0],
        revision_row_indices={9, 10},
    )


# ====================================================================
# main
# ====================================================================
def build() -> Path:
    doc = Document()
    _setup_page(doc)
    _add_page_number_footer(doc)

    add_title_page(doc)
    section_revision_summary(doc)
    section_chapter1(doc)
    section_chapter2(doc)
    section_chapter3(doc)
    section_chapter4(doc)
    section_chapter5(doc)
    section_chapter6(doc)
    section_chapter7(doc)
    section_chapter8(doc)
    section_chapter9(doc)
    section_chapter10(doc)
    section_chapter11(doc)
    section_chapter12(doc)
    section_appendix(doc)
    section_history(doc)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_FILE))
    return OUTPUT_FILE


def main() -> int:
    out = build()
    print(f"[OK] saved: {out}")
    print(f"     size : {out.stat().st_size:,} bytes")
    return 0


if __name__ == "__main__":
    sys.exit(main())
